import io
import os
import tempfile
import re
import json
from typing import Optional

from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import StreamingResponse, Response
from pydantic import BaseModel, field_validator

# Markdown parsing with GitHub-like features
from markdown_it import MarkdownIt
from mdit_py_plugins.tasklists import tasklists_plugin
from mdit_py_plugins.deflist import deflist_plugin
from mdit_py_plugins.footnote import footnote_plugin

# Syntax highlighting
from pygments import highlight
from pygments.lexers import get_lexer_by_name, TextLexer
from pygments.formatters import HtmlFormatter

# PDF rendering - ReportLab (Windows compatible)
REPORTLAB_AVAILABLE = False
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageTemplate, Frame
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY, TA_RIGHT
    from reportlab.lib.colors import black, blue, red, lightgrey, grey
    from reportlab.platypus.doctemplate import BaseDocTemplate
    REPORTLAB_AVAILABLE = True
except Exception as e:
    print(f"ReportLab not available: {e}")

# PDF rendering - WeasyPrint (requires GTK+ libraries)
WEASYPRINT_AVAILABLE = False
try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except Exception as e:
    print(f"WeasyPrint not available: {e}")
    print("WeasyPrint PDF rendering disabled. Using ReportLab instead.")

# DOCX rendering - primary: pandoc
PANDOC_AVAILABLE = False
try:
    import pypandoc
    try:
        # Try to find pandoc; if missing, download it
        pypandoc.get_pandoc_path()
    except OSError:
        pypandoc.download_pandoc()
    PANDOC_AVAILABLE = True
except Exception:
    PANDOC_AVAILABLE = False

# DOCX rendering - fallback: python-docx
PYTHON_DOCX_AVAILABLE = False
try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except Exception:
    PYTHON_DOCX_AVAILABLE = False

# Legacy html2docx import (not used anymore, but kept for compatibility)
HTML2DOCX_AVAILABLE = PYTHON_DOCX_AVAILABLE

# Mermaid diagram rendering
MERMAID_AVAILABLE = False
try:
    import subprocess
    import base64
    from PIL import Image
    MERMAID_AVAILABLE = True
except Exception as e:
    print(f"Mermaid rendering dependencies not available: {e}")
    MERMAID_AVAILABLE = False


app = FastAPI(
    title="Markdown Render API",
    version="1.0.0",
    description="Render Markdown to PDF and DOCX with high fidelity."
)


class RenderRequest(BaseModel):
    markdown: str
    filename: Optional[str] = "document"
    css: Optional[str] = None  # Optional extra CSS for tweaks in PDF/DOCX
    
    @field_validator('markdown')
    @classmethod
    def sanitize_markdown(cls, v):
        """Sanitize control characters from markdown input"""
        if not isinstance(v, str):
            raise ValueError('Markdown must be a string')

        # Remove or replace invalid control characters (except newlines and tabs)
        # Keep \n (0x0A), \r (0x0D), and \t (0x09), remove others
        sanitized = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', v)

        # Normalize line endings
        sanitized = re.sub(r'\r\n|\r', '\n', sanitized)

        return sanitized

    @field_validator('filename')
    @classmethod
    def sanitize_filename(cls, v):
        """Sanitize filename to remove invalid characters"""
        if v is None:
            return "document"

        # Remove invalid filename characters
        sanitized = re.sub(r'[<>:"/\\|?*\x00-\x1F]', '', str(v))

        # Ensure it's not empty after sanitization
        if not sanitized.strip():
            return "document"

        return sanitized.strip()

    @field_validator('css')
    @classmethod
    def sanitize_css(cls, v):
        """Sanitize CSS input"""
        if v is None:
            return None

        # Remove control characters from CSS
        sanitized = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', str(v))
        return sanitized


def _highlight(code: str, lang: Optional[str], attrs):
    try:
        lexer = get_lexer_by_name(lang) if lang else TextLexer()
    except Exception:
        lexer = TextLexer()
    formatter = HtmlFormatter()
    return highlight(code, lexer, formatter)


def _create_md_parser() -> MarkdownIt:
    md = MarkdownIt("gfm-like", {
        "linkify": True,
        "typographer": True,
        "highlight": _highlight
    })
    md = (md
          .use(deflist_plugin)
          .use(tasklists_plugin, enabled=True, label=True, label_after=True)
          .use(footnote_plugin))
    return md


MD = _create_md_parser()

# GitHub-like CSS + Pygments style for highlighted code
BASE_CSS = f"""
:root {{
  --text: #24292f;
  --bg: #ffffff;
  --muted: #57606a;
  --border: #d0d7de;
  --code-bg: #f6f8fa;
}}

html, body {{
  margin: 0; padding: 0; background: var(--bg); color: var(--text);
  font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", "Noto Sans",
               Helvetica, Arial, sans-serif, "Apple Color Emoji", "Segoe UI Emoji";
  line-height: 1.6;
}}

.markdown-body {{
  box-sizing: border-box;
  max-width: 800px;
  margin: 0 auto;
  padding: 24px;
}}

h1, h2, h3, h4, h5, h6 {{
  font-weight: 600; line-height: 1.25; margin: 0.6em 0 0.3em;
}}
h1 {{ font-size: 2em; border-bottom: 1px solid var(--border); padding-bottom: .3em; }}
h2 {{ font-size: 1.5em; border-bottom: 1px solid var(--border); padding-bottom: .3em; }}
h3 {{ font-size: 1.25em; }}

p, ul, ol, dl, blockquote, table, pre {{ margin: 0.4em 0; }}
ul, ol {{ padding-left: 2em; }}

blockquote {{
  padding: 0 1em; color: var(--muted); border-left: .25em solid var(--border);
}}

a {{ color: #0969da; text-decoration: none; }}
a:hover {{ text-decoration: underline; }}

code, kbd, samp {{
  font-family: ui-monospace, SFMono-Regular, Menlo, Consolas,
               "Liberation Mono", monospace; font-size: .95em;
}}

pre {{ background: var(--code-bg); padding: 12px; border-radius: 6px; overflow: auto; }}
pre code {{ background: transparent; border: none; padding: 0; }}

table {{ 
  width: 100%; 
  border-collapse: collapse; 
  table-layout: auto; /* Auto layout for better content fitting */
  word-wrap: break-word; /* Break long words */
  font-size: 0.9em; /* Increased font size for better readability */
  max-width: 100%; /* Ensure table doesn't exceed container */
  margin: 0 auto; /* Center table */
}}
table th, table td {{ 
  border: 1px solid var(--border); 
  padding: 6px 8px; /* Increased padding for better readability */
  word-wrap: break-word;
  overflow-wrap: break-word;
  hyphens: auto; /* Enable hyphenation */
  max-width: 400px; /* Further increased max width for better text distribution */
  white-space: normal; /* Allow text wrapping */
  vertical-align: top; /* Align content to top */
}}
/* Narrow columns for short content (≤3 characters) */
table th:nth-child(n), table td:nth-child(n) {{
  min-width: 20px; /* Minimum width for narrow columns */
}}
/* Specific styling for very narrow content - even smaller widths */
table th:has-text("X"), table td:has-text("X"),
table th:has-text("✓"), table td:has-text("✓"),
table th:has-text("•"), table td:has-text("•") {{
  width: 15px;
  text-align: center;
  padding: 1px;
  min-width: 15px;
  max-width: 15px;
}}

/* Ensure tables use full available width */
table {{
  width: 100% !important;
  max-width: 100% !important;
  table-layout: fixed; /* Better control over column widths */
}}
table th {{ 
  background: #f3f4f6; 
  font-weight: bold;
  font-size: 0.9em;
}}
/* For wide tables, make them even more compact */
table[data-wide="true"] {{
  font-size: 0.7em;
}}
table[data-wide="true"] th, table[data-wide="true"] td {{
  padding: 2px 4px;
  max-width: 200px; /* Smaller max width for wide tables */
}}
/* Ensure tables don't overflow the container */
.markdown-body table {{
  overflow-x: auto;
  display: block;
  white-space: nowrap;
}}
.markdown-body table thead,
.markdown-body table tbody,
.markdown-body table tr {{
  display: table;
  width: 100%;
  table-layout: fixed;
}}
.markdown-body table th,
.markdown-body table td {{
  display: table-cell;
  white-space: normal;
}}

hr {{ border: none; border-top: 1px solid var(--border); margin: 1.5em 0; }}
img {{ max-width: 100%; }}

.task-list-item {{ list-style-type: none; }}
.task-list-item input[type="checkbox"] {{
  margin: 0 .2em .25em -1.4em; vertical-align: middle;
}}

/* Pygments code highlighting */
{HtmlFormatter().get_style_defs('.highlight')}
"""


def markdown_to_html(markdown_text: str, extra_css: Optional[str] = None) -> str:
    # Process Mermaid diagrams before HTML conversion
    processed_markdown = process_mermaid_diagrams_in_markdown(markdown_text)
    
    body_html = MD.render(processed_markdown)
    
    # Post-process HTML to add data-wide attribute to tables with many columns
    body_html = post_process_html_tables(body_html)
    
    css = BASE_CSS + (f"\n/* user CSS overrides */\n{extra_css}" if extra_css else "")
    
    # Add Mermaid CSS and JavaScript
    mermaid_css_js = """
    <script src="https://cdn.jsdelivr.net/npm/mermaid/dist/mermaid.min.js"></script>
    <script>
        mermaid.initialize({ startOnLoad: true, theme: 'default' });
    </script>
    <style>
        .mermaid { text-align: center; margin: 20px 0; }
        .mermaid-placeholder { 
            background: #f8f9fa; 
            border: 2px dashed #dee2e6; 
            padding: 20px; 
            text-align: center; 
            color: #6c757d;
            font-style: italic;
        }
    </style>
    """
    
    html = f"""<!DOCTYPE html>
<html>
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <style>{css}</style>
  {mermaid_css_js}
</head>
<body>
  <article class="markdown-body">
    {body_html}
  </article>
</body>
</html>"""
    return html


def post_process_html_tables(html: str) -> str:
    """Post-process HTML to add data-wide attribute to tables with many columns"""
    import re
    
    # Find all tables in the HTML
    table_pattern = r'<table[^>]*>(.*?)</table>'
    
    def process_table(match):
        table_content = match.group(0)
        table_html = match.group(1)
        
        # Count the number of columns by looking at the first row
        header_pattern = r'<tr[^>]*>(.*?)</tr>'
        header_match = re.search(header_pattern, table_html, re.DOTALL)
        
        if header_match:
            header_content = header_match.group(1)
            # Count <th> or <td> tags
            cell_count = len(re.findall(r'<(th|td)[^>]*>', header_content))
            
            # If table has more than 6 columns, add data-wide attribute
            if cell_count > 6:
                # Add data-wide="true" to the table tag
                table_tag_pattern = r'<table([^>]*)>'
                table_tag_match = re.search(table_tag_pattern, table_content)
                if table_tag_match:
                    attributes = table_tag_match.group(1)
                    new_table_tag = f'<table{attributes} data-wide="true">'
                    return new_table_tag + table_html + '</table>'
        
        return table_content
    
    # Process all tables
    processed_html = re.sub(table_pattern, process_table, html, flags=re.DOTALL)
    
    return processed_html


def process_mermaid_diagrams_in_markdown(markdown_text: str) -> str:
    """Process Mermaid diagrams in markdown and convert them to HTML"""
    lines = markdown_text.split('\n')
    processed_lines = []
    in_code_block = False
    code_lines = []
    code_language = ""

    for line in lines:
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                if code_language and is_mermaid_diagram('\n'.join(code_lines), code_language):
                    # Convert Mermaid to HTML
                    mermaid_code = '\n'.join(code_lines)
                    mermaid_html = f'<div class="mermaid">\n{mermaid_code}\n</div>'
                    processed_lines.append(mermaid_html)
                else:
                    # Regular code block
                    processed_lines.append('```' + code_language)
                    processed_lines.extend(code_lines)
                    processed_lines.append('```')

                code_lines = []
                code_language = ""
                in_code_block = False
            else:
                # Start of code block
                language = line.strip()[3:].strip()
                code_language = language
                in_code_block = True
        elif in_code_block:
            code_lines.append(line)
        else:
            processed_lines.append(line)

    return '\n'.join(processed_lines)


def parse_markdown_table(table_lines):
    """Parse markdown table lines into a ReportLab Table with dynamic width constraints and proper text wrapping"""
    if not table_lines:
        return None

    # Parse table rows
    rows = []
    for line in table_lines:
        # Split by pipe and clean up
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if cells:  # Only add non-empty rows
            # Skip separator rows (rows with only dashes, hyphens, or colons)
            is_separator = True
            for cell in cells:
                # Check if cell contains only dashes, hyphens, colons, or spaces
                if not re.match(r'^[-:\s]+$', cell):
                    is_separator = False
                    break

            if not is_separator:
                rows.append(cells)

    if len(rows) < 2:  # Need at least header and one data row
        return None

    # Calculate available width (A4 page width minus margins)
    page_width = A4[0]  # A4 width in points (595.28 points)
    left_margin = 20  # Minimal margins to maximize table space
    right_margin = 20
    available_width = page_width - left_margin - right_margin  # ~555 points available (maximum space)
    
    num_cols = len(rows[0]) if rows else 1
    
    # Dynamic column width calculation based on content analysis
    col_widths = []
    
    # Analyze content length in each column to determine optimal widths
    max_content_lengths = []
    for col_idx in range(num_cols):
        max_length = 0
        for row in rows:
            if col_idx < len(row):
                # Count characters, considering that long words will wrap
                cell_length = len(row[col_idx])
                max_length = max(max_length, cell_length)
        max_content_lengths.append(max_length)
    
    # Classify columns as narrow or wide with much better thresholds
    narrow_cols = []  # columns with max content ≤ 3 chars (X, ✓, •, etc.)
    medium_cols = []   # columns with content 4-15 chars (short text)
    wide_cols = []     # columns with longer content
    
    for i, content_length in enumerate(max_content_lengths):
        if content_length <= 3:  # Very narrow columns (X markers, etc.)
            narrow_cols.append(i)
        elif content_length <= 15:  # Medium columns (short text)
            medium_cols.append(i)
        else:
            wide_cols.append(i)
    
    # Calculate widths with much better space utilization:
    # - Narrow columns: 15-20 points each (just enough for X + padding)
    # - Medium columns: 40-60 points each (reasonable for short text)
    # - Wide columns: share remaining width proportionally
    
    narrow_width = 15  # Even smaller width for X markers and similar
    medium_width = 45  # Slightly reduced width for short text
    total_narrow_width = len(narrow_cols) * narrow_width
    total_medium_width = len(medium_cols) * medium_width
    remaining_width = available_width - total_narrow_width - total_medium_width
    
    # Ensure we have enough space for wide columns (at least 50% of available width)
    if remaining_width < available_width * 0.5:
        # If too many narrow/medium columns, reduce their width but keep them readable
        if len(narrow_cols) > 0:
            narrow_width = max(12, remaining_width * 0.15 / len(narrow_cols))
        if len(medium_cols) > 0:
            medium_width = max(30, remaining_width * 0.25 / len(medium_cols))
        total_narrow_width = len(narrow_cols) * narrow_width
        total_medium_width = len(medium_cols) * medium_width
        remaining_width = available_width - total_narrow_width - total_medium_width
    
    # Distribute remaining width among wide columns with better proportions
    if wide_cols:
        wide_content_lengths = [max_content_lengths[i] for i in wide_cols]
        total_wide_weight = sum(wide_content_lengths)
        
        for i, col_idx in enumerate(wide_cols):
            if total_wide_weight > 0:
                # Proportional width for wide columns with increased maximum
                proportional_width = (wide_content_lengths[i] / total_wide_weight) * remaining_width
                # Increased maximum width constraint (80% of available width) for better text distribution
                max_width = available_width * 0.8
                col_width = min(max_width, proportional_width)
            else:
                col_width = remaining_width / len(wide_cols)
            col_widths.append((col_idx, col_width))
    else:
        # No wide columns, distribute remaining width among narrow and medium columns
        if len(narrow_cols) > 0 and len(medium_cols) > 0:
            narrow_width = remaining_width * 0.3 / len(narrow_cols)
            medium_width = remaining_width * 0.7 / len(medium_cols)
        elif len(narrow_cols) > 0:
            narrow_width = remaining_width / len(narrow_cols)
        elif len(medium_cols) > 0:
            medium_width = remaining_width / len(medium_cols)
    
    # Add narrow and medium column widths
    for col_idx in narrow_cols:
        col_widths.append((col_idx, narrow_width))
    for col_idx in medium_cols:
        col_widths.append((col_idx, medium_width))
    
    # Sort by column index and extract widths
    col_widths.sort(key=lambda x: x[0])
    col_widths = [width for _, width in col_widths]
    
    # Final safety check - ensure total width doesn't exceed available width
    total_width = sum(col_widths)
    if total_width > available_width:
        # Scale down proportionally
        scale_factor = available_width / total_width
        col_widths = [width * scale_factor for width in col_widths]
    
    # Determine font sizes and padding based on table complexity
    if num_cols > 12:
        font_size_header = 5
        font_size_data = 4
        padding = 0.5
    elif num_cols > 8:
        font_size_header = 6
        font_size_data = 5
        padding = 1
    elif num_cols > 6:
        font_size_header = 7
        font_size_data = 6
        padding = 2
    elif num_cols > 4:
        font_size_header = 8
        font_size_data = 7
        padding = 3
    else:
        font_size_header = 9
        font_size_data = 8
        padding = 4

    # Create paragraph styles for table cells with enhanced text wrapping
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY

    cell_style_header = ParagraphStyle(
        'TableCellHeader',
        fontName='Helvetica-Bold',
        fontSize=font_size_header,
        leading=font_size_header * 1.4,  # Increased leading for better readability
        alignment=TA_LEFT,
        wordWrap='CJK',
        splitLongWords=True,
        keepWithNext=False,
        spaceAfter=0,
        spaceBefore=0
    )

    cell_style_data = ParagraphStyle(
        'TableCellData',
        fontName='Helvetica',
        fontSize=font_size_data,
        leading=font_size_data * 1.4,  # Increased leading for better readability
        alignment=TA_LEFT,
        wordWrap='CJK',
        splitLongWords=True,
        keepWithNext=False,
        spaceAfter=0,
        spaceBefore=0
    )

    # Process table data using Paragraph objects for proper wrapping
    table_data = []
    for i, row in enumerate(rows):
        processed_row = []
        is_header = (i == 0)
        style = cell_style_header if is_header else cell_style_data

        for j, cell in enumerate(row):
            # Process bold text in cells and handle existing <br> tags
            processed_cell = process_bold_text(cell)
            
            # Handle any existing <br> tags in the input text
            processed_cell = processed_cell.replace('<br>', '<br/>')
            processed_cell = processed_cell.replace('<br />', '<br/>')
            
            # For very long content, add line breaks to help with wrapping
            if len(processed_cell) > 80:
                # Insert soft breaks at natural word boundaries
                words = processed_cell.split()
                lines = []
                current_line = ""
                
                for word in words:
                    if len(current_line + " " + word) <= 40:  # Target ~40 chars per line
                        if current_line:
                            current_line += " " + word
                        else:
                            current_line = word
                    else:
                        if current_line:
                            lines.append(current_line)
                            current_line = word
                        else:
                            lines.append(word)
                
                if current_line:
                    lines.append(current_line)
                
                processed_cell = "<br/>".join(lines)
            
            # Create Paragraph object for proper text wrapping
            cell_paragraph = Paragraph(processed_cell, style)
            processed_row.append(cell_paragraph)
        table_data.append(processed_row)

    # Create ReportLab Table with explicit column widths
    table = Table(table_data, colWidths=col_widths, repeatRows=1)

    # Style the table with improved borders and spacing
    table_style = TableStyle([
        # Header row styling
        ('BACKGROUND', (0, 0), (-1, 0), lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, 0), black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),

        # Borders
        ('GRID', (0, 0), (-1, -1), 0.5, black),
        
        # Padding - optimized for readability
        ('LEFTPADDING', (0, 0), (-1, -1), padding),
        ('RIGHTPADDING', (0, 0), (-1, -1), padding),
        ('TOPPADDING', (0, 0), (-1, -1), padding),
        ('BOTTOMPADDING', (0, 0), (-1, -1), padding),
        
        # Enhanced text handling
        ('WORDWRAP', (0, 0), (-1, -1), 'CJK'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
    ])

    table.setStyle(table_style)

    return table


def wrap_text_for_table(text: str, max_length: int = 50) -> str:
    """Wrap long text in table cells to prevent overflow"""
    if len(text) <= max_length:
        return text
    
    # Split text into words
    words = text.split()
    if not words:
        return text
    
    lines = []
    current_line = ""
    
    for word in words:
        # If adding this word would exceed max_length, start a new line
        if len(current_line) + len(word) + 1 > max_length:
            if current_line:
                lines.append(current_line.strip())
                current_line = word
            else:
                # Single word is too long, truncate it
                lines.append(word[:max_length-3] + "...")
                current_line = ""
        else:
            if current_line:
                current_line += " " + word
            else:
                current_line = word
    
    # Add the last line
    if current_line:
        lines.append(current_line.strip())
    
    # Join lines with line breaks
    return "<br/>".join(lines)


def process_bold_text(text: str) -> str:
    """Convert markdown bold syntax to HTML bold tags for ReportLab"""
    import re

    # First, preserve <br/> tags before escaping other HTML
    text = text.replace('<br/>', '___BR_TAG___')
    text = text.replace('<br>', '___BR_TAG___')
    
    # Escape HTML special characters to prevent issues
    # but preserve them for later restoration
    text = text.replace('&', '&amp;')
    text = text.replace('<', '&lt;').replace('>', '&gt;')
    
    # Restore <br/> tags as proper line breaks
    text = text.replace('___BR_TAG___', '<br/>')

    # Handle **bold** syntax
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)

    # Handle __bold__ syntax
    text = re.sub(r'__(.*?)__', r'<b>\1</b>', text)

    return text


def render_mermaid_diagram(mermaid_code: str) -> Optional[bytes]:
    """Render Mermaid diagram to PNG bytes using multiple fallback methods"""
    if not MERMAID_AVAILABLE:
        print("Mermaid rendering not available - missing dependencies")
        return None

    mmd_path = None
    png_path = None

    try:
        # Method 1: Try using mermaid-cli (mmdc) if available
        try:
            # Create temporary files
            with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False) as mmd_file:
                mmd_file.write(mermaid_code)
                mmd_path = mmd_file.name

            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as png_file:
                png_path = png_file.name

            result = subprocess.run([
                'mmdc', '-i', mmd_path, '-o', png_path,
                '--width', '800', '--height', '600',
                '--backgroundColor', 'white'
            ], capture_output=True, text=True, timeout=30)

            if result.returncode == 0 and os.path.exists(png_path) and os.path.getsize(png_path) > 0:
                with open(png_path, 'rb') as f:
                    image_data = f.read()
                print("Mermaid diagram rendered successfully using mmdc")
                return image_data
            else:
                print(f"Mermaid CLI failed: {result.stderr}")
        except (subprocess.TimeoutExpired, FileNotFoundError) as e:
            print(f"Mermaid CLI not available: {e}")

        # Method 2: Try using online Mermaid.ink API (requires internet)
        try:
            import requests
            # Encode mermaid code to base64
            encoded = base64.b64encode(mermaid_code.encode('utf-8')).decode('ascii')
            url = f"https://mermaid.ink/img/{encoded}"

            response = requests.get(url, timeout=15)

            if response.status_code == 200 and len(response.content) > 0:
                print("Mermaid diagram rendered successfully using mermaid.ink")
                return response.content
            else:
                print(f"Online Mermaid API failed with status {response.status_code}")
        except Exception as e:
            print(f"Online Mermaid API failed: {e}")

        # Method 3: Try Kroki API as another fallback
        try:
            import requests
            import zlib
            # Compress and encode
            compressed = zlib.compress(mermaid_code.encode('utf-8'), 9)
            encoded = base64.urlsafe_b64encode(compressed).decode('ascii')
            url = f"https://kroki.io/mermaid/png/{encoded}"

            response = requests.get(url, timeout=15)

            if response.status_code == 200 and len(response.content) > 0:
                print("Mermaid diagram rendered successfully using Kroki")
                return response.content
            else:
                print(f"Kroki API failed with status {response.status_code}")
        except Exception as e:
            print(f"Kroki API failed: {e}")

        return None

    except Exception as e:
        print(f"Error rendering Mermaid diagram: {e}")
        return None
    finally:
        # Clean up temporary files
        try:
            if mmd_path and os.path.exists(mmd_path):
                os.remove(mmd_path)
            if png_path and os.path.exists(png_path):
                os.remove(png_path)
        except Exception:
            pass


def is_mermaid_diagram(code_block: str, language: str) -> bool:
    """Check if a code block is a Mermaid diagram"""
    return language and language.lower() in ['mermaid', 'mmd']


def validate_docx_file(docx_bytes: bytes) -> bool:
    """Validate that the bytes represent a valid DOCX file"""
    try:
        # Check if it's empty
        if len(docx_bytes) == 0:
            return False
        
        # Check ZIP signature (DOCX files are ZIP archives)
        if not docx_bytes.startswith(b'PK'):
            return False
        
        # Try to parse as ZIP to ensure it's valid
        import zipfile
        with zipfile.ZipFile(io.BytesIO(docx_bytes), 'r') as zip_file:
            # Check for required DOCX files
            required_files = ['[Content_Types].xml', 'word/document.xml']
            file_list = zip_file.namelist()
            
            for required_file in required_files:
                if required_file not in file_list:
                    return False
        
        return True
    except Exception as e:
        print(f"DOCX validation failed: {e}")
        return False


def create_mermaid_placeholder(mermaid_code: str) -> str:
    """Create a text placeholder for Mermaid diagrams when rendering fails"""
    lines = mermaid_code.strip().split('\n')
    first_line = lines[0] if lines else ""
    
    # Extract diagram type from first line
    diagram_type = "Diagram"
    if first_line.startswith('graph'):
        diagram_type = "Flowchart"
    elif first_line.startswith('sequenceDiagram'):
        diagram_type = "Sequence Diagram"
    elif first_line.startswith('classDiagram'):
        diagram_type = "Class Diagram"
    elif first_line.startswith('stateDiagram'):
        diagram_type = "State Diagram"
    elif first_line.startswith('erDiagram'):
        diagram_type = "Entity Relationship Diagram"
    elif first_line.startswith('journey'):
        diagram_type = "User Journey"
    elif first_line.startswith('gantt'):
        diagram_type = "Gantt Chart"
    elif first_line.startswith('pie'):
        diagram_type = "Pie Chart"
    
    return f"[{diagram_type} - Mermaid diagram rendering not available]"


class NumberedCanvas(canvas.Canvas):
    """Custom canvas with page numbers and headers/footers"""
    
    def __init__(self, *args, **kwargs):
        canvas.Canvas.__init__(self, *args, **kwargs)
        self._saved_page_states = []
    
    def draw_page_number(self, page_num, document_title="Document"):
        """Draw page number and header/footer"""
        try:
            # Get page dimensions safely
            if hasattr(self, '_pagesize') and self._pagesize:
                page_width, page_height = self._pagesize
            else:
                # Fallback to A4 size
                page_width, page_height = A4
            
            # Header - Leave blank as requested (no document title)
            # Just add a subtle line for visual separation - FULL PAGE WIDTH
            self.setStrokeColor(grey)
            self.setLineWidth(0.5)
            self.line(0, page_height - 30, page_width, page_height - 30)
            
            # Footer - Generation date at bottom left (aligned with full-width line)
            from datetime import datetime
            current_date = datetime.now().strftime("%d-%b-%Y")
            self.setFont("Helvetica", 9)
            self.setFillColor(grey)
            self.drawString(30, 30, current_date)  # Reduced margin to align with full-width line
            
            # Footer - "Confidential" at bottom center
            self.setFont("Helvetica-Bold", 9)
            self.setFillColor(black)
            confidential_text = "Confidential"
            confidential_width = self.stringWidth(confidential_text, "Helvetica-Bold", 9)
            self.drawString((page_width - confidential_width) / 2, 30, confidential_text)
            
            # Footer - Page number at bottom right (aligned with full-width line)
            self.setFont("Helvetica", 9)
            self.setFillColor(grey)
            page_text = f"Page {page_num}"
            text_width = self.stringWidth(page_text, "Helvetica", 9)
            self.drawString(page_width - text_width - 30, 30, page_text)  # Reduced margin to align with full-width line
            
            # Footer line - FULL PAGE WIDTH
            self.setStrokeColor(grey)
            self.setLineWidth(0.5)
            self.line(0, 50, page_width, 50)
        except Exception as e:
            # If pagination fails, just continue without it
            print(f"Warning: Could not draw page numbers: {e}")
            pass
    
    def showPage(self):
        """Override showPage to add page numbers"""
        try:
            # Get document title from the template if available
            document_title = getattr(self, '_document_title', 'Document')
            self.draw_page_number(self._pageNumber, document_title)
        except Exception as e:
            print(f"Warning: Could not add page numbers: {e}")
        finally:
            canvas.Canvas.showPage(self)


class NumberedDocTemplate(SimpleDocTemplate):
    """Custom document template with numbered pages using SimpleDocTemplate"""
    
    def __init__(self, filename, document_title="Document", **kwargs):
        self.document_title = document_title
        SimpleDocTemplate.__init__(self, filename, **kwargs)
    
    def build(self, flowables, onFirstPage=None, onLaterPages=None, canvasmaker=NumberedCanvas):
        """Build the document with custom canvas"""
        # Create a custom canvas maker that knows about the document title
        def custom_canvas_maker(*args, **kwargs):
            canvas = canvasmaker(*args, **kwargs)
            canvas._document_title = self.document_title
            return canvas
        
        SimpleDocTemplate.build(self, flowables, canvasmaker=custom_canvas_maker)


def markdown_to_pdf_bytes_reportlab(markdown_text: str, extra_css: Optional[str] = None) -> bytes:
    """Convert markdown to PDF using ReportLab with enhanced markdown support and pagination"""
    if not REPORTLAB_AVAILABLE:
        raise RuntimeError("ReportLab is not available for PDF generation")
    
    # Use generic document title instead of extracting from markdown
    # This prevents "Protocol Title:" from appearing in headers
    document_title = "Document"
    
    # Create a BytesIO buffer to hold the PDF
    buffer = io.BytesIO()
    
    # Create PDF document with better margins and space for headers/footers
    try:
        doc = NumberedDocTemplate(buffer, document_title=document_title, pagesize=A4, 
                                rightMargin=20, leftMargin=20, 
                                topMargin=30, bottomMargin=50)
    except Exception as e:
        print(f"Warning: Custom template failed, using SimpleDocTemplate: {e}")
        # Fallback to SimpleDocTemplate without pagination
        doc = SimpleDocTemplate(buffer, pagesize=A4, 
                               rightMargin=20, leftMargin=20, 
                               topMargin=30, bottomMargin=50)
    
    # Get styles
    styles = getSampleStyleSheet()
    
    # Create enhanced custom styles with reduced spacing
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=6,
        spaceBefore=0,
        alignment=TA_LEFT,
        textColor=black,
        fontName='Helvetica-Bold',
        leading=22
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=4,
        spaceBefore=6,
        alignment=TA_LEFT,
        textColor=black,
        fontName='Helvetica-Bold',
        leading=18
    )

    subheading_style = ParagraphStyle(
        'CustomSubHeading',
        parent=styles['Heading3'],
        fontSize=12,
        spaceAfter=3,
        spaceBefore=4,
        alignment=TA_LEFT,
        textColor=black,
        fontName='Helvetica-Bold',
        leading=16
    )

    subsubheading_style = ParagraphStyle(
        'CustomSubSubHeading',
        parent=styles['Heading4'],
        fontSize=11,
        spaceAfter=2,
        spaceBefore=3,
        alignment=TA_LEFT,
        textColor=black,
        fontName='Helvetica-Bold',
        leading=14
    )

    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=3,
        spaceBefore=1,
        alignment=TA_LEFT,
        textColor=black,
        fontName='Helvetica',
        leading=14,
        wordWrap='CJK'
    )
    
    bullet_style = ParagraphStyle(
        'CustomBullet',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=3,
        spaceBefore=2,
        alignment=TA_LEFT,
        textColor=black,
        fontName='Helvetica',
        leftIndent=20,
        bulletIndent=10,
        leading=14,
        wordWrap='CJK'
    )
    
    code_style = ParagraphStyle(
        'CustomCode',
        parent=styles['Code'],
        fontSize=8,
        spaceAfter=6,
        spaceBefore=6,
        alignment=TA_LEFT,
        textColor=black,
        fontName='Courier',
        backColor='#f5f5f5',
        borderColor='#cccccc',
        borderWidth=1,
        borderPadding=8,
        leading=11,
        wordWrap='CJK'
    )
    
    # Parse markdown and convert to PDF elements
    elements = []
    lines = markdown_text.split('\n')
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    
    code_language = ""

    for i, line in enumerate(lines):
        original_line = line
        line = line.strip()

        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                if code_lines:
                    code_text = '\n'.join(code_lines)

                    # Check if this is a Mermaid diagram
                    if code_language and is_mermaid_diagram(code_text, code_language):
                        # Try to render Mermaid diagram
                        image_data = render_mermaid_diagram(code_text)
                        if image_data:
                            try:
                                # Add image to PDF with proper aspect ratio
                                from reportlab.platypus import Image as RLImage
                                from PIL import Image as PILImage

                                # Save image data to BytesIO and pass it directly
                                img_buffer = io.BytesIO(image_data)

                                # Get original image dimensions to maintain aspect ratio
                                pil_img = PILImage.open(io.BytesIO(image_data))
                                orig_width, orig_height = pil_img.size

                                # Calculate dimensions to fit within page width while maintaining aspect ratio
                                max_width = 450  # Max width in points (about 6.25 inches)
                                aspect_ratio = orig_height / orig_width

                                # Scale to fit max width
                                new_width = min(max_width, orig_width)
                                new_height = new_width * aspect_ratio

                                img_buffer.seek(0)  # Reset buffer position
                                img = RLImage(img_buffer, width=new_width, height=new_height)
                                elements.append(img)
                                elements.append(Spacer(1, 12))
                            except Exception as e:
                                print(f"Error adding Mermaid image to PDF: {e}")
                                # Fallback to placeholder text
                                placeholder = create_mermaid_placeholder(code_text)
                                elements.append(Paragraph(placeholder, code_style))
                        else:
                            # Fallback to placeholder text
                            placeholder = create_mermaid_placeholder(code_text)
                            elements.append(Paragraph(placeholder, code_style))
                    else:
                        # Regular code block
                        elements.append(Paragraph(code_text, code_style))

                    code_lines = []
                    code_language = ""
                in_code_block = False
            else:
                # Start of code block - extract language
                code_language = line[3:].strip()
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(original_line)
            continue
        
        # Handle tables
        if '|' in line and not line.startswith('#'):
            # This looks like a table row
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(original_line)
            continue
        elif in_table:
            # End of table
            if table_lines:
                table = parse_markdown_table(table_lines)
                if table:
                    elements.append(table)
                    elements.append(Spacer(1, 12))  # Add space after table
            in_table = False
            table_lines = []
        
        # Handle empty lines
        if not line:
            elements.append(Spacer(1, 12))
            continue
        
        # Handle headers with better detection
        if line.startswith('#### '):
            text = line[5:].strip()
            elements.append(Paragraph(text, subsubheading_style))
        elif line.startswith('### '):
            text = line[4:].strip()
            elements.append(Paragraph(text, subheading_style))
        elif line.startswith('## '):
            text = line[3:].strip()
            elements.append(Paragraph(text, heading_style))
        elif line.startswith('# '):
            text = line[2:].strip()
            elements.append(Paragraph(text, title_style))
        
        # Handle bullet points with better detection
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # Process bold text
            text = process_bold_text(text)
            # Handle any existing <br> tags in the input text
            text = text.replace('<br>', '<br/>')
            text = text.replace('<br />', '<br/>')
            # Handle nested bullets
            indent_level = 0
            while original_line.startswith('  '):
                indent_level += 1
                original_line = original_line[2:]
            if indent_level > 0:
                bullet_text = '  ' * indent_level + '• ' + text
            else:
                bullet_text = '• ' + text
            elements.append(Paragraph(bullet_text, bullet_style))
        
        # Handle numbered lists
        elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
            processed_line = process_bold_text(line)
            # Handle any existing <br> tags in the input text
            processed_line = processed_line.replace('<br>', '<br/>')
            processed_line = processed_line.replace('<br />', '<br/>')
            elements.append(Paragraph(processed_line, normal_style))
        
        # Handle regular text (plain text only)
        else:
            if line:
                processed_line = process_bold_text(line)
                # Handle any existing <br> tags in the input text
                processed_line = processed_line.replace('<br>', '<br/>')
                processed_line = processed_line.replace('<br />', '<br/>')
                elements.append(Paragraph(processed_line, normal_style))
    
    # Handle any remaining code block
    if code_lines:
        code_text = '\n'.join(code_lines)
        elements.append(Paragraph(code_text, code_style))
    
    # Handle any remaining table
    if table_lines:
        table = parse_markdown_table(table_lines)
        if table:
            elements.append(table)
            elements.append(Spacer(1, 12))
    
    # Build PDF
    try:
        if isinstance(doc, NumberedDocTemplate):
            doc.build(elements, canvasmaker=NumberedCanvas)
        else:
            doc.build(elements)
    except Exception as e:
        print(f"Warning: Custom build failed, trying standard build: {e}")
        # Fallback to standard build
        doc.build(elements)
    
    # Get PDF bytes
    pdf_bytes = buffer.getvalue()
    buffer.close()
    
    return pdf_bytes


def markdown_to_pdf_bytes(markdown_text: str, extra_css: Optional[str] = None) -> bytes:
    """Convert markdown to PDF with fallback to ReportLab if WeasyPrint unavailable"""
    # Try WeasyPrint first (better HTML/CSS support)
    if WEASYPRINT_AVAILABLE:
        html = markdown_to_html(markdown_text, extra_css=extra_css)
        pdf_bytes = HTML(string=html, base_url=".").write_pdf(stylesheets=[CSS(string=BASE_CSS)])
        return pdf_bytes
    
    # Fallback to ReportLab (Windows compatible)
    if REPORTLAB_AVAILABLE:
        return markdown_to_pdf_bytes_reportlab(markdown_text, extra_css)
    
    # No PDF backend available
    raise RuntimeError(
        "PDF rendering is not available. Neither WeasyPrint nor ReportLab is installed. "
        "Please install ReportLab (pip install reportlab) or use DOCX rendering instead."
    )


def _create_reference_docx() -> str:
    """Create a reference DOCX file with headers and footers for pagination"""
    if not HTML2DOCX_AVAILABLE:
        # If we can't create a reference doc, return empty string
        print("HTML2DOCX not available, cannot create reference document")
        return ""
    
    # Create a temporary reference document
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_ref:
            tmp_ref_path = tmp_ref.name
        
        # Create a document with headers and footers
        doc = Document()
        
        # Add a sample paragraph to ensure the document has content
        doc.add_paragraph("Reference Document for Headers and Footers")
        
        # Get the sections
        sections = doc.sections
        
        for section in sections:
            # Set header
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = "Confidential"
            header_para.alignment = 0  # Left alignment
            # Make header bold
            try:
                from docx.shared import Pt
                header_run = header_para.runs[0] if header_para.runs else header_para.add_run()
                header_run.bold = True
                header_run.font.size = Pt(10)
            except Exception as e:
                print(f"Could not style header in reference doc: {e}")
            
            # Set footer with page number
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = 1  # Center alignment

            # Add page number field
            try:
                from docx.oxml.shared import qn
                from docx.oxml import OxmlElement

                # Clear any existing content
                footer_para.clear()

                # Add "Page " text as a run
                footer_para.add_run("Page ")

                # Create a new run for the page number field
                run_element = OxmlElement('w:r')

                # Create page number field within the run
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                run_element.append(fldChar1)

                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = "PAGE"
                run_element.append(instrText)

                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'end')
                run_element.append(fldChar2)

                # Add the field run to the footer paragraph
                footer_para._p.append(run_element)
            except Exception as e:
                print(f"Could not add page number field to reference doc: {e}")
                # Fallback to static text
                footer_para.clear()
                footer_para.text = "Page [Page Number]"
        
        # Save the reference document
        doc.save(tmp_ref_path)
        print(f"Created reference document: {tmp_ref_path}")
        return tmp_ref_path
        
    except Exception as e:
        print(f"Error creating reference DOCX: {e}")
        return ""


def markdown_to_docx_bytes(markdown_text: str, extra_css: Optional[str] = None) -> bytes:
    """
    Prefer Pandoc for best fidelity with pagination. Fallback to HTML->DOCX with html2docx.
    Enhanced with better error handling and file validation.
    """
    # Handle empty or whitespace-only markdown
    if not markdown_text or not markdown_text.strip():
        markdown_text = "# Empty Document\n\nThis document contains no content."
    
    if PANDOC_AVAILABLE:
        # Pandoc needs an output file path for binary outputs
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_out:
            tmp_out_path = tmp_out.name
        try:
            # Try to create reference document for pagination
            reference_doc_path = _create_reference_docx()
            
            # Prepare extra args
            extra_args = ["--standalone"]
            
            # Only add reference document if it was created successfully
            if reference_doc_path and os.path.exists(reference_doc_path):
                extra_args.append(f"--reference-doc={reference_doc_path}")
                print(f"Using reference document: {reference_doc_path}")
            else:
                print("No reference document available, generating DOCX without custom pagination")
            
            # Convert with Pandoc
            try:
                pypandoc.convert_text(
                    markdown_text,
                    to="docx",
                    format="md",
                    outputfile=tmp_out_path,
                    extra_args=extra_args
                )
                
                # Validate the generated file
                if not os.path.exists(tmp_out_path):
                    raise RuntimeError("Pandoc failed to generate output file")
                
                file_size = os.path.getsize(tmp_out_path)
                if file_size == 0:
                    raise RuntimeError("Pandoc generated empty file")
                
                # Additional validation: try to read the file
                with open(tmp_out_path, "rb") as f:
                    data = f.read()
                
                # Validate DOCX file structure using comprehensive validation
                if not validate_docx_file(data):
                    raise RuntimeError("Generated file is not a valid DOCX file")
                
                print(f"Pandoc successfully generated DOCX file ({file_size} bytes)")
                return data
            except Exception as e:
                print(f"Pandoc conversion failed: {e}")
                # If Pandoc fails, fall through to HTML2DOCX fallback
                pass
        finally:
            try:
                if os.path.exists(tmp_out_path):
                    os.remove(tmp_out_path)
                # Clean up reference document if it was created
                if 'reference_doc_path' in locals() and reference_doc_path and os.path.exists(reference_doc_path):
                    os.remove(reference_doc_path)
            except Exception as e:
                print(f"Warning: Could not clean up temporary files: {e}")
                pass

    if HTML2DOCX_AVAILABLE:
        # Enhanced DOCX generation with proper structure
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        # Create document
        doc = Document()
        
        # Set up sections with headers and footers
        section = doc.sections[0]
        
        # Add header
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = "Confidential"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if header_para.runs:
            header_para.runs[0].font.bold = True
            header_para.runs[0].font.size = Pt(10)
        
        # Add footer with page numbers
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add page number field
        try:
            from datetime import datetime
            current_date = datetime.now().strftime("%d-%b-%Y")

            # Clear any existing content
            footer_para.clear()

            # Add date and "Page " text as a run
            run = footer_para.add_run(f"{current_date} | Page ")

            # Create a new run for the page number field
            run_element = OxmlElement('w:r')

            # Add page number field within the run
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run_element.append(fldChar1)

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = "PAGE"
            run_element.append(instrText)

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run_element.append(fldChar2)

            # Append the run with the field to the paragraph
            footer_para._p.append(run_element)
        except Exception as e:
            print(f"Could not add page numbers: {e}")
            footer_para.clear()
            footer_para.text = "Confidential"

        # Parse markdown with proper formatting
        lines = markdown_text.split('\n')
        in_table = False
        table_lines = []
        in_list = False
        
        for line in lines:
            stripped = line.strip()
            
            # Handle empty lines
            if not stripped:
                if in_table:
                    # Process accumulated table
                    if len(table_lines) >= 2:
                        _add_table_to_doc(doc, table_lines)
                    in_table = False
                    table_lines = []
                in_list = False
                continue
            
            # Handle tables
            if '|' in stripped and not stripped.startswith('#'):
                if not in_table:
                    in_table = True
                    table_lines = []
                table_lines.append(line)
                continue
            elif in_table:
                # End of table
                if len(table_lines) >= 2:
                    _add_table_to_doc(doc, table_lines)
                in_table = False
                table_lines = []
            
            # Handle headings - use explicit black color
            if stripped.startswith('# '):
                heading = doc.add_heading(stripped[2:], level=1)
                for run in heading.runs:
                    from docx.shared import RGBColor
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Explicit black
            elif stripped.startswith('## '):
                heading = doc.add_heading(stripped[3:], level=2)
                for run in heading.runs:
                    from docx.shared import RGBColor
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Explicit black
            elif stripped.startswith('### '):
                heading = doc.add_heading(stripped[4:], level=3)
                for run in heading.runs:
                    from docx.shared import RGBColor
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Explicit black
            elif stripped.startswith('#### '):
                heading = doc.add_heading(stripped[5:], level=4)
                for run in heading.runs:
                    from docx.shared import RGBColor
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Explicit black
            elif stripped.startswith('##### '):
                heading = doc.add_heading(stripped[6:], level=5)
                for run in heading.runs:
                    from docx.shared import RGBColor
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Explicit black
            elif stripped.startswith('###### '):
                heading = doc.add_heading(stripped[7:], level=6)
                for run in heading.runs:
                    from docx.shared import RGBColor
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Explicit black
            
            # Handle bullet points
            elif stripped.startswith(('- ', '* ')):
                text = stripped[2:]
                para = doc.add_paragraph(text, style='List Bullet')
                _apply_inline_formatting(para)
                in_list = True
            
            # Handle numbered lists
            elif stripped[:2].isdigit() and stripped[2:4] in ('. ', ') '):
                text = stripped[stripped.index(' ')+1:]
                para = doc.add_paragraph(text, style='List Number')
                _apply_inline_formatting(para)
                in_list = True
            
            # Handle horizontal rules
            elif stripped in ('---', '***', '___'):
                doc.add_paragraph('_' * 50)
            
            # Regular paragraphs
            else:
                para = doc.add_paragraph(stripped)
                _apply_inline_formatting(para)
        
        # Handle any remaining table
        if in_table and len(table_lines) >= 2:
            _add_table_to_doc(doc, table_lines)

        # Save document
        buf = io.BytesIO()
        try:
            doc.save(buf)
            buf.seek(0)
            data = buf.getvalue()
            
            # Validate
            if len(data) == 0:
                raise RuntimeError("Generated empty DOCX file")
            
            if not data.startswith(b'PK'):
                raise RuntimeError("Generated file is not a valid DOCX")
            
            print(f"Successfully generated DOCX file ({len(data)} bytes)")
            return data
            
        except Exception as e:
            print(f"Error saving DOCX: {e}")
            raise RuntimeError(f"Failed to save DOCX: {str(e)}")
        finally:
            buf.close()

    raise RuntimeError(
        "No DOCX backend available. Install pypandoc (recommended) or python-docx."
    )


def _process_cell_markdown(cell, text, is_header=False):
    """Process markdown formatting in table cell text"""
    from docx.shared import RGBColor
    import re

    # Clear the cell's default paragraph
    cell.text = ''

    # Split by <br> tags for line breaks
    lines = re.split(r'<br\s*/?>', text, flags=re.IGNORECASE)

    for line_idx, line in enumerate(lines):
        # Add paragraph for each line (except first which uses existing paragraph)
        if line_idx == 0:
            para = cell.paragraphs[0]
        else:
            para = cell.add_paragraph()

        # Parse **bold** syntax
        parts = re.split(r'(\*\*.*?\*\*)', line)

        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = para.add_run(part[2:-2])
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                # Regular text
                run = para.add_run(part)
                if is_header:
                    run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)


def _add_table_to_doc(doc, table_lines):
    """Helper function to add a table to the document with markdown formatting support"""
    from docx.shared import RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Parse table
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if cells and not all(re.match(r'^[-:\s]+$', cell) for cell in cells):
            rows.append(cells)

    if len(rows) < 2:
        return

    # Create table with plain grid style
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'  # Plain black grid style

    # Fill table with markdown processing
    for i, row_data in enumerate(rows):
        row_cells = table.rows[i].cells
        for j, cell_text in enumerate(row_data):
            if j < len(row_cells):
                # Process markdown in cell text
                _process_cell_markdown(row_cells[j], cell_text, is_header=(i == 0))


def _apply_inline_formatting(para):
    """Apply bold and other inline formatting to a paragraph"""
    text = para.text
    para.clear()
    
    # Handle **bold** syntax
    import re
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.font.bold = True
        else:
            para.add_run(part)


@app.post("/test-docx")
async def test_docx():
    """Test endpoint to generate a simple DOCX file"""
    try:
        # Create a simple test document
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Add simple content
        doc.add_heading("Test Document", level=1)
        doc.add_paragraph("This is a test document to verify DOCX generation works.")
        doc.add_paragraph("If you can see this, the DOCX file is working correctly.")
        
        # Add footer
        from datetime import datetime
        current_date = datetime.now().strftime("%d-%b-%Y")
        footer_para = doc.add_paragraph(f"{current_date} | Test Document")
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save to buffer
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        data = buf.getvalue()
        buf.close()
        
        # Validate
        if len(data) == 0 or not data.startswith(b'PK'):
            raise RuntimeError("Generated invalid DOCX file")
        
        return Response(
            content=data,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=test_document.docx"}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Test DOCX generation failed: {str(e)}")


@app.post("/render/docx-raw")
async def render_docx_raw(request: Request):
    """Render DOCX from raw JSON body with enhanced error handling"""
    try:
        # Get raw body and parse JSON
        body = await request.body()
        data = json.loads(body.decode('utf-8'))
        
        # Extract fields
        markdown = data.get('markdown', '')
        filename = data.get('filename', 'document')
        css = data.get('css', None)
        
        # Validate required fields
        if not markdown:
            raise HTTPException(status_code=400, detail="markdown field is required")
        
        # Sanitize filename
        filename = re.sub(r'[<>:"/\\|?*\x00-\x1F]', '', str(filename)).strip()
        if not filename:
            filename = "document"
        
        # Generate DOCX with validation
        docx_bytes = markdown_to_docx_bytes(markdown, extra_css=css)
        
        # Final validation before sending
        if not validate_docx_file(docx_bytes):
            raise HTTPException(status_code=500, detail="Generated DOCX file is corrupted")
        
        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}.docx"',
                "Content-Length": str(len(docx_bytes))
            }
        )
        
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=400, detail=f"Invalid JSON: {str(e)}")
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e))
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except Exception as e:
        print(f"Unexpected error in DOCX rendering: {e}")
        raise HTTPException(status_code=500, detail=f"DOCX rendering error: {str(e)}")


@app.post("/render/pdf-raw")
async def render_pdf_raw(request: Request):
    """Render PDF from raw JSON body"""
    try:
        # Get raw body and parse JSON
        body = await request.body()
        data = json.loads(body.decode('utf-8'))
        
        # Extract fields
        markdown = data.get('markdown', '')
        filename = data.get('filename', 'document')
        css = data.get('css', None)
        
        # Validate required fields
        if not markdown:
            raise HTTPException(status_code=400, detail="markdown field is required")
        
        # Generate PDF
        pdf_bytes = markdown_to_pdf_bytes(markdown, extra_css=css)
        
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}.pdf"'
            }
        )
        
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=400, detail=f"Invalid JSON: {str(e)}")
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF rendering error: {str(e)}")


@app.post("/render/pdf")
def render_pdf(req: RenderRequest):
    try:
        pdf_bytes = markdown_to_pdf_bytes(req.markdown, extra_css=req.css)
        filename = req.filename.replace('"', "").replace("\n", "")
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}.pdf"'
            }
        )
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF rendering error: {str(e)}")


@app.post("/render/docx")
def render_docx(req: RenderRequest):
    """Render DOCX with enhanced error handling and validation"""
    try:
        docx_bytes = markdown_to_docx_bytes(req.markdown, extra_css=req.css)
        
        # Final validation before sending
        if not validate_docx_file(docx_bytes):
            raise HTTPException(status_code=500, detail="Generated DOCX file is corrupted")
        
        filename = req.filename.replace('"', "").replace("\n", "")
        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}.docx"',
                "Content-Length": str(len(docx_bytes))
            }
        )
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e))
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except Exception as e:
        print(f"Unexpected error in DOCX rendering: {e}")
        raise HTTPException(status_code=500, detail=f"DOCX rendering error: {str(e)}")


@app.get("/")
def root():
    return {
        "message": "Markdown Render API",
        "version": "1.0.0",
        "endpoints": {
            "POST /render/pdf": "Convert markdown to PDF (standard)",
            "POST /render/pdf-raw": "Convert markdown to PDF (handles control characters)",
            "POST /render/docx": "Convert markdown to DOCX (standard)",
            "POST /render/docx-raw": "Convert markdown to DOCX (handles control characters)",
            "GET /health": "Health check"
        },
        "request_format": {
            "markdown": "string (required)",
            "filename": "string (optional, defaults to 'document')",
            "css": "string (optional, extra CSS for styling)"
        },
        "features": {
            "pdf_rendering": WEASYPRINT_AVAILABLE,
            "docx_rendering": PANDOC_AVAILABLE or HTML2DOCX_AVAILABLE,
            "pandoc_available": PANDOC_AVAILABLE,
            "html2docx_available": HTML2DOCX_AVAILABLE,
            "mermaid_rendering": MERMAID_AVAILABLE
        }
    }


@app.get("/health")
def health_check():
    return {
        "status": "healthy",
        "features": {
            "pdf_rendering": WEASYPRINT_AVAILABLE or REPORTLAB_AVAILABLE,
            "weasyprint_available": WEASYPRINT_AVAILABLE,
            "reportlab_available": REPORTLAB_AVAILABLE,
            "docx_rendering": PANDOC_AVAILABLE or HTML2DOCX_AVAILABLE,
            "pandoc_available": PANDOC_AVAILABLE,
            "html2docx_available": HTML2DOCX_AVAILABLE,
            "mermaid_rendering": MERMAID_AVAILABLE
        }
    }
